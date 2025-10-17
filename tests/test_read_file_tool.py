"""Tests for the read_file MCP tool."""

from __future__ import annotations

import importlib.util
import io
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
from types import ModuleType
from typing import ClassVar
from unittest import TestCase
from unittest.mock import MagicMock, patch
from contextlib import redirect_stderr


def _ensure_stub_dependencies() -> None:
    """Provide lightweight stand-ins for optional dependencies during tests."""
    if "youtube_transcript_api" not in sys.modules:
        stub = ModuleType("youtube_transcript_api")

        class _StubYouTubeTranscriptApi:
            def fetch(self, video_id: str) -> list[dict[str, object]]:
                # Minimal stub: tests never call this path, but returning an empty
                # transcript keeps the server logic happy if it does.
                return []

        stub.YouTubeTranscriptApi = _StubYouTubeTranscriptApi  # type: ignore[attr-defined]
        sys.modules["youtube_transcript_api"] = stub


def _load_mcp_module() -> ModuleType:
    """Load the content_mcp module directly from its path for test usage."""
    module_path = (
        Path(__file__).resolve().parents[1]
        / "mcp-server"
        / "content_mcp.py"
    )
    spec = importlib.util.spec_from_file_location("content_mcp", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError("Unable to load content_mcp module for testing")
    _ensure_stub_dependencies()
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _write_simple_pdf(path: Path, text: str) -> None:
    """Create a minimal PDF file containing the provided text."""
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT\n/F1 12 Tf\n72 720 Td\n({escaped}) Tj\nET\n".encode("utf-8")

    catalog = b"<< /Type /Catalog /Pages 2 0 R >>\n"
    pages = b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>\n"
    page = (
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\n"
    )
    stream = (
        b"<< /Length "
        + str(len(content)).encode("ascii")
        + b" >>\nstream\n"
        + content
        + b"endstream\n"
    )
    font = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"

    objects = [catalog, pages, page, stream, font]

    pdf = bytearray()
    pdf.extend(b"%PDF-1.4\n")
    offsets = [0]

    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        if not obj.endswith(b"\n"):
            pdf.extend(b"\n")
        pdf.extend(b"endobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for start in offsets[1:]:
        pdf.extend(f"{start:010} 00000 n \n".encode("ascii"))

    pdf.extend(
        (
            f"trailer<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )

    path.write_bytes(pdf)


class ReadFileToolTests(TestCase):
    content_mcp: ClassVar[ModuleType]

    @classmethod
    def setUpClass(cls) -> None:
        cls.content_mcp = _load_mcp_module()

    def setUp(self) -> None:
        self._temp_dir = TemporaryDirectory()
        self.temp_path = Path(self._temp_dir.name)

    def tearDown(self) -> None:
        self._temp_dir.cleanup()

    def test_read_text_file_returns_contents(self) -> None:
        file_path = self.temp_path / "sample.txt"
        file_path.write_text("Line 1\nLine 2", encoding="utf-8")

        result = self.content_mcp.read_file(str(file_path))

        self.assertEqual(result, "Line 1\nLine 2")

    def test_read_markdown_truncates_output(self) -> None:
        file_path = self.temp_path / "notes.md"
        file_path.write_text("A" * 120, encoding="utf-8")

        result = self.content_mcp.read_file(str(file_path), max_chars=50)

        self.assertTrue(result.endswith("...[truncated]"))
        self.assertLessEqual(len(result), 50 + len("\n\n...[truncated]"))

    def test_read_docx_file_extracts_text(self) -> None:
        from docx import Document  # type: ignore[import-not-found]

        doc_path = self.temp_path / "document.docx"
        document = Document()
        document.add_heading("Sample Heading", level=1)
        document.add_paragraph("Body paragraph content.")
        document.save(str(doc_path))

        result = self.content_mcp.read_file(str(doc_path))

        self.assertIn("Sample Heading", result)
        self.assertIn("Body paragraph content.", result)

    def test_read_pdf_file_extracts_text(self) -> None:
        pdf_path = self.temp_path / "document.pdf"
        _write_simple_pdf(pdf_path, "PDF body text.")

        result = self.content_mcp.read_file(str(pdf_path))

        self.assertIn("PDF body text.", result)

    def test_read_doc_file_uses_best_effort_fallback(self) -> None:
        doc_path = self.temp_path / "legacy.doc"
        payload = b"\xd0\xcf\x11\xe0" + b"Hello DOC content!" + b"\x00" * 8
        doc_path.write_bytes(payload)

        with patch.object(self.content_mcp, "_extract_doc_with_textract", return_value=None):
            result = self.content_mcp.read_file(str(doc_path))

        self.assertIn("Hello DOC content", result)

    def test_read_file_rejects_unsupported_extension(self) -> None:
        data_path = self.temp_path / "data.json"
        data_path.write_text("{}", encoding="utf-8")

        with self.assertRaises(ValueError):
            self.content_mcp.read_file(str(data_path))

    def test_read_file_rejects_directory_path(self) -> None:
        with self.assertRaises(ValueError):
            self.content_mcp.read_file(str(self.temp_path))

    def test_read_file_debug_logging_emits_messages(self) -> None:
        buffer = io.StringIO()
        with patch.object(self.content_mcp, "_READ_FILE_DEBUG", True):
            with patch.object(self.content_mcp, "_LOGGER") as mock_logger:
                mock_logger.handlers = []
                mock_logger.setLevel = MagicMock()
                mock_logger.info = MagicMock()
                mock_logger.addHandler = MagicMock()
                mock_logger.propagate = True

                with self.assertRaises(FileNotFoundError):
                    self.content_mcp.read_file("/definitely/not/here.txt")

        mock_logger.addHandler.assert_called()
        mock_logger.setLevel.assert_called()
        mock_logger.info.assert_called()
