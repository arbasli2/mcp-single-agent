from pathlib import Path
import importlib
import logging
import os
import re
import socket
import sys
from html.parser import HTMLParser
from urllib import error as urlerror
from urllib import parse as urlparse
from urllib import request as urlrequest

from mcp.server.fastmcp import FastMCP
from youtube_transcript_api import YouTubeTranscriptApi

# Configure logging to file
def _setup_logging() -> None:
    """Setup logging configuration to write to file instead of stderr."""
    script_dir = os.path.dirname(__file__)
    log_file = os.path.join(script_dir, "mcp-server.log")
    
    # Create rotating file handler to prevent log files from getting too large
    from logging.handlers import RotatingFileHandler
    handler = RotatingFileHandler(
        log_file, 
        maxBytes=10*1024*1024,  # 10MB max file size
        backupCount=5  # Keep 5 backup files
    )
    
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    handler.setFormatter(formatter)
    
    # Configure root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    # Clear any existing handlers to avoid duplicates
    root_logger.handlers.clear()
    root_logger.addHandler(handler)
    
    # Set specific logger levels
    logging.getLogger("mcp").setLevel(logging.WARNING)
    logging.getLogger("read_file").setLevel(logging.INFO)
    
    # Log the startup
    logging.info("="*50)
    logging.info("MCP Content Server logging initialized")
    logging.info(f"Log file: {log_file}")
    logging.info("="*50)

# Setup logging
_setup_logging()

_LOGGER = logging.getLogger("read_file")
_READ_FILE_DEBUG = os.getenv("READ_FILE_DEBUG", "false").lower() == "true"


def _debug_log(message: str) -> None:
    if _READ_FILE_DEBUG:
        _LOGGER.info(f"[debug] {message}")
    else:
        _LOGGER.debug(f"[debug] {message}")

# Create an MCP server
mcp = FastMCP("content-mcp")


# Create prompt
@mcp.prompt()
def system_prompt() -> str:
    """Instructions for the multi-purpose content agent"""
    script_dir = os.path.dirname(__file__)
    prompt_path = os.path.join(script_dir, "prompts", "system_instructions.md")
    with open(prompt_path, "r") as file:
        return file.read()


# Create tool
@mcp.tool()
def fetch_video_transcript(url: str) -> str:
    """
    Extract transcript with timestamps from a YouTube video URL and format it for analysis.

    Use this tool when:
    - User provides a YouTube URL and wants video content analysis
    - User asks to summarize, analyze, or process a YouTube video
    - User wants to create content based on a YouTube video

    Args:
        url (str): YouTube video URL (youtube.com or youtu.be format)

    Returns:
        str: Formatted transcript with timestamps, where each entry is on a new line
             in the format: "[MM:SS] Text"
    """
    logging.info(f"Fetching video transcript for URL: {url}")
    
    # Extract video ID from URL
    video_id_pattern = r'(?:v=|\/)([0-9A-Za-z_-]{11}).*'
    video_id_match = re.search(video_id_pattern, url)

    if not video_id_match:
        raise ValueError("Invalid YouTube URL")

    video_id = video_id_match.group(1)

    try:
        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)

        # Format each entry with timestamp and text
        formatted_entries = []
        for entry in transcript:
            # Convert seconds to MM:SS format
            minutes = int(entry.start // 60)
            seconds = int(entry.start % 60)
            timestamp = f"[{minutes:02d}:{seconds:02d}]"

            formatted_entry = f"{timestamp} {entry.text}"
            formatted_entries.append(formatted_entry)

        # Join all entries with newlines
        result = "\n".join(formatted_entries)
        logging.info(f"Successfully fetched transcript with {len(formatted_entries)} entries")
        return result

    except Exception as e:
        logging.error(f"Error fetching transcript for {url}: {str(e)}")
        raise Exception(f"Error fetching transcript: {str(e)}")


@mcp.tool()
def fetch_instructions(prompt_name: str) -> str:
    """
    Fetch specialized writing instructions and guidelines for creating different types of content.

    Use this tool when user wants to:
    - Write a blog post (use prompt_name: "write_blog_post")
    - Create social media content (use prompt_name: "write_social_post")
    - Generate video chapter timestamps (use prompt_name: "write_video_chapters")

    Args:
        prompt_name (str): Type of instructions to fetch
        Available options:
            - "write_blog_post": Guidelines for writing blog posts
            - "write_social_post": Guidelines for social media content
            - "write_video_chapters": Guidelines for creating video chapter timestamps

    Returns:
        str: Detailed instructions and guidelines for the requested content type
    """
    logging.info(f"Fetching instructions for prompt: {prompt_name}")
    
    script_dir = os.path.dirname(__file__)
    prompt_path = os.path.join(script_dir, "prompts", f"{prompt_name}.md")
    with open(prompt_path, "r") as f:
        return f.read()


# Helper functions for file reading
def _read_text_file(path: Path) -> str:
    _debug_log(f"Reading text/markdown file: {path}")
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        _debug_log(f"Successfully read {len(content)} characters from text file")
        return content
    except Exception as e:
        logging.error(f"Failed to read text file {path}: {str(e)}")
        raise RuntimeError(f"Cannot read text file: {str(e)}") from e


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        _debug_log("python-docx not installed when reading docx")
        raise RuntimeError(
            "Reading .docx files requires the optional dependency 'python-docx'."
        ) from exc

    _debug_log(f"Reading docx file via python-docx: {path}")
    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    # Include simple table extraction so important content is not missed.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _read_doc_file(path: Path) -> str:
    _debug_log(f"Attempting to read legacy doc file: {path}")
    textract_text = _extract_doc_with_textract(path)
    if textract_text:
        return textract_text

    fallback_text = _extract_doc_basic(path)
    if fallback_text:
        return fallback_text

    _debug_log("Failed to extract .doc content with both textract and fallback")
    raise RuntimeError(
        "Unable to extract text from .doc file. Install the 'textract' package for improved support."
    )


def _extract_doc_with_textract(path: Path) -> str | None:
    try:
        textract = importlib.import_module("textract")
    except ImportError:
        _debug_log("textract not available for .doc extraction")
        return None

    try:
        text_bytes = textract.process(str(path))
    except Exception:
        _debug_log("textract process call failed; falling back")
        return None

    text = text_bytes.decode("utf-8", errors="replace").strip()
    return text or None


def _extract_doc_basic(path: Path) -> str | None:
    _debug_log("Using best-effort fallback for .doc extraction")
    raw = path.read_bytes()
    if not raw:
        return None

    ascii_chunks = re.findall(rb"[\x09\x0A\x0D\x20-\x7E]{4,}", raw)
    decoded_lines: list[str] = []
    seen: set[str] = set()
    for chunk in ascii_chunks:
        cleaned_bytes = re.sub(rb"\s+", b" ", chunk).strip()
        if not cleaned_bytes:
            continue
        text = cleaned_bytes.decode("latin-1", errors="ignore").replace("\x00", "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        decoded_lines.append(text)

    candidate = "\n".join(decoded_lines).strip()
    if candidate:
        return candidate

    unicode_candidate = raw.decode("utf-16-le", errors="ignore").replace("\x00", "").strip()
    return unicode_candidate or None


def _read_pdf_file(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text  # type: ignore[import-not-found]
    except ImportError as exc:
        _debug_log("pdfminer.six missing when reading PDF")
        raise RuntimeError(
            "Reading PDF files requires the optional dependency 'pdfminer.six'."
        ) from exc

    try:
        text = extract_text(str(path))
    except Exception as exc:
        _debug_log(f"pdfminer failed to extract text: {exc}")
        raise RuntimeError(f"Failed to extract text from PDF: {exc}") from exc

    return text


def _normalise_text(text: str) -> str:
    normalised = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalised.splitlines()]
    cleaned_lines = [line for line in lines if line]
    return "\n".join(cleaned_lines).strip()


@mcp.tool()
def read_file(file_path: str, max_chars: int = 6000) -> str:
    """Read textual content from supported file types.

    Supported extensions:
        - .txt (plain text)
        - .md / .markdown (Markdown)
        - .doc (legacy Microsoft Word)
        - .docx (Word Open XML)
        - .pdf (Portable Document Format)

    Args:
        file_path (str): Absolute or relative path to the file on disk.
        max_chars (int): Optional upper bound on returned characters.

    Returns:
        str: Extracted text content trimmed to ``max_chars`` characters.
    """
    logging.info(f"Reading file: {file_path} (max_chars: {max_chars})")
    
    _debug_log(f"read_file invoked with file_path={file_path!r}, max_chars={max_chars}")

    resolved_path = Path(file_path).expanduser()
    try:
        resolved_path = resolved_path.resolve(strict=True)
    except FileNotFoundError as exc:
        _debug_log(f"File not found during resolve: {resolved_path}")
        logging.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}") from exc

    if not resolved_path.is_file():
        _debug_log(f"Resolved path is not a file: {resolved_path}")
        logging.error(f"Path does not point to a file: {resolved_path}")
        raise ValueError(f"Path does not point to a file: {resolved_path}")

    try:
        limit = int(max_chars)
    except (TypeError, ValueError) as exc:
        raise ValueError("max_chars must be an integer") from exc

    limit = max(1, min(limit, 200_000))
    _debug_log(f"Resolved path={resolved_path}, suffix={resolved_path.suffix.lower()}, limit={limit}")

    suffix = resolved_path.suffix.lower()
    try:
        if suffix in {".txt", ".md", ".markdown"}:
            raw_text = _read_text_file(resolved_path)
        elif suffix == ".docx":
            raw_text = _read_docx_file(resolved_path)
        elif suffix == ".doc":
            raw_text = _read_doc_file(resolved_path)
        elif suffix == ".pdf":
            raw_text = _read_pdf_file(resolved_path)
        else:
            _debug_log(f"Unsupported suffix encountered: {suffix}")
            logging.error(f"Unsupported file type '{suffix}' for file: {file_path}")
            raise ValueError(
                "Unsupported file type. Supported extensions: .txt, .md, .markdown, .doc, .docx, .pdf"
            )
    except Exception as e:
        logging.error(f"Error reading file {file_path}: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to read file {file_path}: {str(e)}") from e

    cleaned = _normalise_text(raw_text)
    if not cleaned:
        _debug_log("No text extracted after normalisation")
        logging.warning(f"No text content could be extracted from file: {file_path}")
        return "No textual content could be extracted from the file."

    if len(cleaned) <= limit:
        _debug_log(f"Returning full content length={len(cleaned)}")
        logging.info(f"Successfully read file {file_path} ({len(cleaned)} chars)")
        return cleaned

    _debug_log(f"Content truncated from length={len(cleaned)} to limit={limit}")
    logging.info(f"Successfully read file {file_path} (truncated from {len(cleaned)} to {limit} chars)")
    return f"{cleaned[:limit]}\n\n...[truncated]"


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._tokens: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        elif tag in {"p", "div", "section", "article", "li", "br", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._tokens.append("\n")

    def handle_endtag(self, tag):
        if tag in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1
        elif tag in {"p", "div", "section", "article", "li"}:
            self._tokens.append("\n")

    def handle_data(self, data):
        if self._skip_depth > 0:
            return
        stripped = data.strip()
        if stripped:
            self._tokens.append(stripped)

    def get_text(self) -> str:
        lines: list[str] = []
        current: list[str] = []
        for token in self._tokens:
            if token == "\n":
                if current:
                    lines.append(" ".join(current).strip())
                    current = []
            else:
                current.append(token)
        if current:
            lines.append(" ".join(current).strip())
        return "\n".join(line for line in lines if line)


@mcp.tool()
def fetch_web_content(url: str, max_chars: int = 6000, timeout_seconds: int = 20) -> str:
    """Fetch readable text content from a public webpage.

    Args:
        url (str): HTTP or HTTPS URL to fetch.
        max_chars (int): Optional upper bound on returned characters to prevent huge responses.
        timeout_seconds (int): Maximum seconds to wait for the HTTP request.

    Returns:
        str: Extracted body text, truncated to ``max_chars`` characters when needed.
    """
    logging.info(f"Fetching web content from: {url} (max_chars: {max_chars}, timeout: {timeout_seconds}s)")
    
    parsed = urlparse.urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("Only http and https URLs are supported")

    timeout = max(1, min(int(timeout_seconds), 120))
    req = urlrequest.Request(url, headers={"User-Agent": "Mozilla/5.0 (compatible; content-mcp-agent/1.0)"})
    try:
        with urlrequest.urlopen(req, timeout=timeout) as response:
            charset = response.headers.get_content_charset() or "utf-8"
            html_bytes = response.read()
    except urlerror.URLError as exc:
        reason = exc.reason
        if isinstance(reason, socket.timeout):
            error_msg = f"Timed out after {timeout} seconds while fetching the URL"
            logging.error(f"Timeout fetching {url}: {error_msg}")
            raise Exception(error_msg)
        error_msg = f"Error fetching URL: {reason}"
        logging.error(f"URL error for {url}: {error_msg}")
        raise Exception(error_msg)

    try:
        html_text = html_bytes.decode(charset, errors="replace")
    except LookupError:
        html_text = html_bytes.decode("utf-8", errors="replace")

    extractor = _HTMLTextExtractor()
    extractor.feed(html_text)
    text = extractor.get_text()

    if not text:
        logging.warning(f"No readable text content found at {url}")
        return "No readable text content found at the provided URL."

    trimmed = text[:max_chars]
    if len(text) > max_chars:
        trimmed += "\n\n...[truncated]"
        logging.info(f"Successfully fetched web content from {url} (truncated from {len(text)} to {max_chars} chars)")
    else:
        logging.info(f"Successfully fetched web content from {url} ({len(text)} chars)")
    return trimmed


if __name__ == "__main__":
    logging.info("Starting MCP content server...")
    try:
        mcp.run(transport="stdio")
    except Exception as e:
        logging.error(f"MCP server error: {e}", exc_info=True)
        raise



